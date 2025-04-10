# This file contains helper functions for performing model selection associated with Figure 4\

import glob
import os
import json
import re

import pandas as pd
import numpy as np
import seaborn as sns
from scipy.optimize import minimize
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn

import matplotlib.pyplot as plt
import statsmodels.formula.api as smf
from statsmodels.stats.outliers_influence import variance_inflation_factor
from statsmodels.discrete.count_model import ZeroInflatedNegativeBinomialP
from statsmodels.api import NegativeBinomial
import statsmodels.api as sm
from statsmodels.formula.api import ols
from statsmodels.genmod.families.links import logit
from statsmodels.genmod.families import Binomial
from statsmodels.regression.mixed_linear_model import MixedLM

# Suppress FutureWarning about BIC calculation change
from statsmodels.genmod.generalized_linear_model import SET_USE_BIC_LLF
SET_USE_BIC_LLF(True)

def generate_model_list(predictors, response_vars):
	"""
	Generate a list of models based on predictors and response variables.

	Args:
	predictors (list of str): List of predictor variable names.
	response_vars (list of str): List of response variable names.

	Returns:
	list of tuples: Each tuple contains a model formula and the response variable.
	"""
	model_list = []

	for response_var in response_vars:
		for predictor in predictors:
			
			
			if predictor not in ['season']:

				model_list.append((f'{response_var} ~ season*{predictor}'
								   , response_var))
			else:
				model_list.append((f'{response_var} ~ {predictor}', response_var))

	return model_list

def check_multicollinearity(df, predictors):
	"""Check for multicollinearity among predictors."""
	# Ensure predictors are numeric
	df_predictors = df[predictors].select_dtypes(include=[np.number])

	# Fill or drop missing values; here we'll drop rows with any NaNs to avoid bias from imputation
	df_predictors = df_predictors.dropna()

	# Check for infinite values and remove them
	df_predictors = df_predictors.replace([np.inf, -np.inf], np.nan).dropna()

	if df_predictors.empty:
		raise ValueError("No valid numeric predictors available after cleaning.")

	vif_data = pd.DataFrame()
	vif_data["feature"] = df_predictors.columns
	vif_data["VIF"] = [variance_inflation_factor(df_predictors.values, i) for i in range(len(df_predictors.columns))]

	high_vif = vif_data[vif_data["VIF"] > 5]
	if not high_vif.empty:
		print("Warning: High multicollinearity detected for the following predictors:\n", high_vif)
	else:
		print("No high multicollinearity detected.")

	return vif_data

def fit_zinb_models(df, models):
	"""Fit Zero-Inflated Negative Binomial models and return AIC values with relevant information."""

	results = []
	models_dict = {}

	for formula, response in models:
		try:
			# Fit the Zero-Inflated Negative Binomial model
			response_var = response.strip()
			predictor_terms = formula.split('~')[1].strip().split('+')
			predictors = extract_predictors(predictor_terms)
			response_and_predictors = [response_var] + predictors

			df_clean = df.replace([np.inf, -np.inf], np.nan)
			df_clean = df_clean.dropna(subset=response_and_predictors)
			df_clean = df_clean.drop_duplicates(subset=response_and_predictors)
			
			model = ZeroInflatedNegativeBinomialP.from_formula(formula, data=df_clean).fit(method='bfgs', disp=False, maxiter=100)
			
			# Calculate pseudo R-squared
			llf = model.llf 
			llnull = model.llnull  
			pseudo_r2 = 1 - (llf / llnull)
			
			results.append({
				'Model': formula,
				'Response': response,
				'AIC': model.aic,
				'BIC': model.bic,
				'Pseudo R-squared': pseudo_r2,
				'Converged': model.mle_retvals.get('converged', False),
				'Iterations': model.mle_retvals.get('iterations', 'N/A')
			})
			models_dict[formula] = model
			
		except Exception as e:
			print(f"Error fitting model {formula}: {e}")

	return pd.DataFrame(results), models_dict

def fit_mixed_effects_models(df, models, random_effect):
	"""Fit mixed-effects models with random_effect as the random effect and return AIC values with relevant information."""

	results = []
	models_dict = {}

	for formula, response in models:
		try:
			# Debug prints to check data
			#print(f"Fitting model: {formula}")
			#print(f"Data shape: {df.shape}")
			#print(f"Columns used in model: {formula}")
			response_and_predictors = [response.strip()] + formula.split('~')[1].strip().split('+')
			df_clean = df.dropna(subset=response_and_predictors)
			#print(f"Cleaned data shape (after dropping NAs): {df_clean.shape}")

			# Specify the mixed model with 'deployment' as the random effect
			model = smf.mixedlm(formula, data=df_clean, groups=df[random_effects]).fit(disp=False)
			results.append({'Model': formula, 'Response': response, 'AIC': model.aic, 'BIC': model.bic})
			models_dict[formula] = model
		except Exception as e:
			print(f"Error fitting model {formula}: {e}")

	return pd.DataFrame(results), models_dict

def fit_glm_nb_models(df, models, random_effect=None):
	"""Fit Negative Binomial models and optionally include a random effect."""

	results = []
	fitted_models = {}

	for formula, response in models:
		try:
			# Fit the Negative Binomial model using GLM or Mixed Effects Model
			print(f"Fitting model: {formula}")
			response_var = response.strip()
			predictor_terms = formula.split('~')[1].strip().split('+')

			# Extract individual predictors considering interaction terms
			predictors = extract_predictors(predictor_terms)
			response_and_predictors = [response_var] + predictors

			df_clean = df.dropna(subset=response_and_predictors).drop_duplicates(subset=response_and_predictors)

			if random_effect:

				# Extract response and predictors
				y, X = smf.ols(formula, data=df_clean).endog, smf.ols(formula, data=df_clean).exog
				group = df_clean[random_effect]

				# Fit a Poisson mixed-effect model
				mixed_model = MixedLM(y, X, groups=group).fit()

				# Optimize alpha for Negative Binomial
				alpha_init = 1.0
				result = minimize(neg_llf, alpha_init, args=(y, mixed_model.fittedvalues), bounds=[(0.0001, 10)])
				alpha_opt = result.x[0]
				#print(f"\tOptimized dispersion parameter (alpha): {alpha_opt}")

				# Fit the GLM with Negative Binomial family and optimized alpha
				glmm_model = sm.GLM(y, mixed_model.fittedvalues, family=sm.families.NegativeBinomial(alpha=alpha_opt)).fit()
				model = glmm_model

			else:

				# Extract response and predictors
				y, X = smf.ols(formula, data=df_clean).endog, smf.ols(formula, data=df_clean).exog

				# Initial guess and bounds for alpha
				alpha_init = 1.0
				bounds = [(0.0001, 10)]  # Avoid alpha approaching 0 to prevent numerical issues

				# Minimize neg_llf to find the optimal alpha
				result = minimize(neg_llf, alpha_init, args=(y, X), bounds=bounds)

				# Get the optimal alpha
				alpha_opt = result.x[0]
				#print(f"\tOptimized dispersion parameter (alpha): {alpha_opt}")

				# Fit the model using GLM with Negative Binomial family and optimized alpha
				glm_model = smf.glm(formula=formula, data=df_clean, family=sm.families.NegativeBinomial(alpha=alpha_opt)).fit()
				model = glm_model


			# Calculate pseudo R-squared
			results_as_html = model.summary().tables[0].as_html()
			df_summary = pd.read_html(results_as_html, header=0, index_col=0)[0]
			df_summary = df_summary.set_index('No. Observations:')
			df_summary = df_summary.drop(columns = response_var)
			pseudo_r2 = df_summary.loc['Pseudo R-squ. (CS):'][0]

			results.append({
				'Model': formula,
				'Response': response,
				'AIC': model.aic,
				'BIC': model.bic,
				'Pseudo R-squared': pseudo_r2,
				'Converged': model.mle_retvals.get('converged', False) if hasattr(model, 'mle_retvals') else True,
			})
			fitted_models[formula] = model
		except Exception as e:
			print(f"Error fitting model {formula}: {e}")
			results.append({
				'Model': formula,
				'Response': response,
				'AIC': None,
				'BIC': None,
				'Pseudo R-squared': None,
				'Converged': False,
				'Error': str(e)
			})

	return pd.DataFrame(results), fitted_models


def fit_beta_models(df, models):
    """Fit Beta-like regression models using GLM with Binomial family and logit link, and return AIC values with relevant information."""

    results = []
    fitted_models = {}

    for formula, response in models:
        try:
            # Fit the Beta-like regression model
            print(f"Fitting model: {formula}")
            response_and_predictors = [response.strip()] + formula.split('~')[1].strip().split('+')
            df_clean = df.dropna(subset=response_and_predictors)
            model = smf.glm(formula=formula, data=df_clean, family=Binomial(link=logit())).fit()
            results.append({
                'Model': formula,
                'Response': response,
                'AIC': model.aic,
                'BIC': model.bic,
                'Params': model.params,
                'P-values': model.pvalues
            })
            fitted_models[formula] = model
        except Exception as e:
            print(f"Error fitting model {formula}: {e}")

    return pd.DataFrame(results), fitted_models

def fit_quasibinomial_models(df, models):
    """Fit Quasi-Binomial regression models and return AIC values with relevant information."""

    results = []
    fitted_models = {}

    for formula, response in models:
        try:
            # Fit the Quasi-Binomial regression model
            print(f"Fitting model: {formula}")
            response_and_predictors = [response.strip()] + formula.split('~')[1].strip().split('+')
            df_clean = df.dropna(subset=response_and_predictors)
            model = smf.glm(formula=formula, data=df_clean, family=Binomial(link=logit())).fit(scale='X2')
            results.append({
                'Model': formula,
                'Response': response,
                'AIC': model.aic,
                'BIC': model.bic,
                'Params': model.params,
                'P-values': model.pvalues
            })
            fitted_models[formula] = model
        except Exception as e:
            print(f"Error fitting model {formula}: {e}")
            results.append({
                'Model': formula,
                'Response': response,
                'AIC': None,
                'BIC': None,
                'Converged': False,
                'Iterations': 'N/A',
                'Warnings': str(e),
                'Params': None,
                'P-values': None
            })

    return pd.DataFrame(results), fitted_models

def fit_ols_models(df, models):
	"""Fit Ordinary Least Squares models and return AIC values with relevant information."""

	results = []
	models_dict = {}

	for formula, response in models:
		try:
			# Fit the Ordinary Least Squares model
			response_and_predictors = [response.strip()] + formula.split('~')[1].strip().split('+')
			df_clean = df.dropna(subset=response_and_predictors)
			model = ols(formula, data=df_clean).fit()
			results.append({
				'Model': formula,
				'Response': response,
				'AIC': model.aic,
				'BIC': model.bic
			})
			models_dict[formula] = model
		except Exception as e:
			print(f"Error fitting model {formula}: {e}")

	return pd.DataFrame(results), models_dict

def best_model_assessment(df_results, criterion):
	"""Assess and print the best model."""
	grouped = df_results.groupby('Response')
	for name, group in grouped:
		if criterion == 'Pseudo R-squared':
			best_model = group.loc[group[criterion].idxmax()]
		else:
			best_model = group.loc[group[criterion].idxmin()]
		print(f"Best model for {name}:")
		print(f"Model: {best_model['Model']}")
		print(f"AIC: {best_model['AIC']}")
		print(f"BIC: {best_model['BIC']}")
		print(f"pseudo-R-squared: {best_model['Pseudo R-squared']}\n")


def clean_data(df):
    # Replace infinite values with NaN
    df.replace([np.inf, -np.inf], np.nan, inplace=True)
    # Drop rows with NaN values (includes the ones that were infinite)
    df.dropna(inplace=True)
    return df


def check_numeric_columns(df, predictors):
    numeric_predictors = df[predictors].select_dtypes(include=[np.number])
    return numeric_predictors.columns.tolist()

def visualize_residuals(model, title):
    """Visualize the residuals of the model."""
    residuals = model.resid
    fitted = model.fittedvalues
    
    fig, ax = plt.subplots(1, 2, figsize=(14, 6))
    sns.histplot(residuals, kde=True, ax=ax[0])
    ax[0].set_title('Residuals Distribution - ' + title)
    
    sm.qqplot(residuals, line='45', ax=ax[1])
    ax[1].set_title('Q-Q Plot - ' + title)
    
    plt.show()
	
def extract_predictors(predictor_terms):
    """Extract individual predictors considering interaction terms."""
    predictors = []
    for term in predictor_terms:
        term = term.strip()
        if '*' in term:
            parts = term.split('*')
            predictors.extend([p.strip() for p in parts])
        elif ':' in term:
            parts = term.split(':')
            predictors.extend([p.strip() for p in parts])
        else:
            predictors.append(term)
    return list(set(predictors))


def neg_llf(alpha, x, intercept):
	"""
	From https://github.com/statsmodels/statsmodels/issues/9031
	"""

	try:
		model = sm.GLM(x, intercept, family=sm.families.NegativeBinomial(alpha=alpha)).fit()
		return -model.llf
	except:
		return np.inf
	
	

def save_ANOVA_tables_to_word(formula, anova_table, tukey_result, file_path):
	"""
	Saves ANOVA and Tukey HSD results to a Word document.

	Parameters:
	- anova_table: DataFrame containing ANOVA results.
	- tukey_result: Tukey HSD results from statsmodels MultiComparison.
	- file_path: Path to save the Word document.
	"""
	# Create a Word Document
	doc = Document()
	apply_font_styles(doc)

	if not 'supplemental' in formula:
		figure = formula.split('figure')[1].split('_')[0]
		panel = formula.split(':')[0].split('panel')[-1]
		model_text = formula.split(':')[-1]
		new_formula = (' ').join(['Figure', figure, 'panel', panel+':', model_text])
		analysis_heading = doc.add_heading(f"{new_formula}", level=2)
		customize_heading(analysis_heading, font_name='Arial', font_size=9, color=(0, 0, 0))
	else:
		figure = formula.split('figure')[1].split('_')[0]
		panel = formula.split(':')[0].split('panel')[-1]
		model_text = formula.split(':')[-1]
		new_formula = (' ').join(['Supplemental Figure', figure, 'panel', panel+':', model_text])
		analysis_heading = doc.add_heading(f"{new_formula}", level=2)
		customize_heading(analysis_heading, font_name='Arial', font_size=9, color=(0, 0, 0))

	# Add ANOVA Table
	anova_heading = doc.add_heading('ANOVA', level=2)
	customize_heading(anova_heading, font_name='Arial', font_size=6, color=(128, 128, 128))

	# Extract columns and data from the DataFrame
	anova_columns = anova_table.reset_index().columns.tolist()  # Ensure index is included if it serves as a data column
	anova_data = anova_table.reset_index().values.tolist()      # Get all values including those from the index

	anova_table_doc = doc.add_table(rows=1, cols=len(anova_columns))
	hdr_cells = anova_table_doc.rows[0].cells
	for idx, name in enumerate(anova_columns):
		hdr_cells[idx].text = name

	for row in anova_data:
		row_cells = anova_table_doc.add_row().cells
		for idx, item in enumerate(row):
			row_cells[idx].text = f"{item:.3f}" if isinstance(item, float) else str(item)

	set_table_font(anova_table_doc)
	style_table_header(anova_table_doc)
	set_table_borders(anova_table_doc)
	set_row_height(anova_table_doc)

	if len(tukey_result) > 1:
		# Add Tukey HSD Table
		tukey_heading = doc.add_heading('Tukey HSD', level=2)
		customize_heading(tukey_heading, font_name='Arial', font_size=6, color=(128, 128, 128))

		# Extract columns and data from the DataFrame
		tukey_columns = tukey_result.columns.tolist()
		tukey_data = tukey_result.values

		tukey_table_doc = doc.add_table(rows=1, cols=len(tukey_columns))
		hdr_cells = tukey_table_doc.rows[0].cells
		for idx, name in enumerate(tukey_columns):
			hdr_cells[idx].text = name

		for row in tukey_data:
			row_cells = tukey_table_doc.add_row().cells
			for idx, item in enumerate(row):
				row_cells[idx].text = str(item)

		set_table_font(tukey_table_doc)
		style_table_header(tukey_table_doc)
		set_table_borders(tukey_table_doc)
		set_row_height(tukey_table_doc)
		
	# Save the document
	doc.save(file_path)
	print(f"Results saved to {file_path}")
	

def save_glm_tables_to_word(formula, tables, file_path='GLM_Results.docx'):
	"""
	Saves multiple summary tables to a Word document.

	Parameters:
	- formula: The formula used in the model.
	- tables: A list of DataFrames representing the summary tables.
	- file_path: The path where the Word document will be saved.
	"""
	doc = Document()

	# Add each table to the document
	for i, table_df in enumerate(tables, start=1):
		# Add a heading for context
		if i == 1:
			analysis_heading = doc.add_heading(f'GLM results: {formula}', level=1)
			customize_heading(analysis_heading, font_name='Arial', font_size=11, color=(0, 0, 0))

			num_rows, num_cols = table_df.shape
			table_master = doc.add_table(rows=num_rows, cols=num_cols)

			for row_idx in range(num_rows):
				row_cells = table_master.rows[row_idx].cells
				for col_idx in range(num_cols):
					item = table_df.iat[row_idx, col_idx]
					row_cells[col_idx].text = str(item)

					# Make the text in the first and third columns bold
					if col_idx == 0 or col_idx == 2:
						for run in row_cells[col_idx].paragraphs[0].runs:
							run.font.bold = True

			# Set bottom border of the last row to be double
			last_row = table_master.rows[-1]
			for cell in last_row.cells:
				set_bottom_border_double(cell)

			set_row_height(table_master, height=200)
			set_table_font(table_master, font_name='Arial', font_size=6)
			set_table_borders(table_master)
		else:  # Handle the second table
			# Include index and column names for the second table
			table_df = table_df.reset_index().rename(columns = {"index":""})  # Include index as a column
			columns = table_df.columns.tolist()

			num_rows, num_cols = table_df.shape
			table_master = doc.add_table(rows=num_rows + 1, cols=len(columns))
			hdr_cells = table_master.rows[0].cells

			# Set headers for table in Word
			for idx, col_name in enumerate(columns):
				hdr_cells[idx].text = str(col_name)

			# Populate table with data
			for row_idx, row in enumerate(table_df.itertuples(index=False), start=1):
				row_cells = table_master.rows[row_idx].cells
				row_values = list(row)
				for col_idx, item in enumerate(row_values):
					row_cells[col_idx].text = f"{item:.3f}" if isinstance(item, float) else str(item)

			# Style header row
			for cell in hdr_cells:
				cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
				for run in cell.paragraphs[0].runs:
					run.font.bold = True

			set_row_height(table_master, height=200)
			set_table_font(table_master, font_name='Arial', font_size=6)
			set_table_borders(table_master)

	# Save the document
	doc.save(file_path)
	print(f"Results saved to {file_path}")

def set_bottom_border_double(cell):
	"""Sets a double bottom border for a cell."""
	tc = cell._element
	tc_pr = tc.get_or_add_tcPr()
	tc_borders = tc_pr.find(qn('w:tcBorders'))
	if tc_borders is None:
		tc_borders = OxmlElement('w:tcBorders')
		tc_pr.append(tc_borders)

	bottom = OxmlElement('w:bottom')
	bottom.set(qn('w:val'), 'double')
	bottom.set(qn('w:sz'), '6')  # Adjust size as needed
	bottom.set(qn('w:space'), '0')
	bottom.set(qn('w:color'), 'auto')
	tc_borders.append(bottom)
	
def apply_font_styles(doc):
    """Apply universal font styles (Arial, 11pt) to the entire document."""
    # Set the default font for 'Normal' style
    normal_style = doc.styles['Normal']
    font = normal_style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Set default font for heading styles
    for heading_number in range(1, 4):  # e.g., Heading1, Heading2, Heading3
        style_name = f'Heading {heading_number}'
        if style_name in doc.styles:
            style = doc.styles[style_name]
            font = style.font
            font.name = 'Arial'
            font.size = Pt(11)

def set_table_font(table, font_name='Arial', font_size=6):
    """Sets the font for all text in a table and centers text in each cell."""
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # Vertical alignment
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Horizontal alignment
                for run in paragraph.runs:
                    run.font.name = font_name
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                    run.font.size = Pt(font_size)
					
def customize_heading(paragraph, font_name='Arial', font_size=11, color=(0, 0, 0)):
	"""
	Customizes the font and color of a given paragraph.

	Parameters:
	- paragraph: The paragraph object to customize.
	- font_name: The name of the font to use.
	- font_size: The size of the font.
	- color: A tuple with RGB values for the text color.
	"""
	run = paragraph.runs[0]
	run.font.name = font_name
	run.font.size = Pt(font_size)
	run.font.color.rgb = RGBColor(*color)
	run.font.bold = False
	
def style_table_header(table):
	"""Style the first row as bold and add a single bottom border."""
	hdr_row = table.rows[0]
	for cell in hdr_row.cells:
		for paragraph in cell.paragraphs:
			for run in paragraph.runs:
				run.font.bold = True
				run.font.size = Pt(6)
				# Underline can be set like this
				run.font.underline = False

	# Set a single border for the bottom of the header row
	for cell in hdr_row.cells:
		cell_xml = cell._element
		tcPr = cell_xml.get_or_add_tcPr()
		tcBorders = parse_xml("<w:tcBorders xmlns:w='http://schemas.openxmlformats.org/wordprocessingml/2006/main'>"
							  "<w:bottom w:val='double' w:sz='4' w:space='0' w:color='auto'/>"
							  "</w:tcBorders>")
		tcPr.append(tcBorders)
		

def set_table_borders(table):
    tbl = table._tbl
    tblBorders = OxmlElement('w:tblBorders')
    
    border_attrs = {
        'w:val': 'single',
        'w:sz': '4',  # border size
        'w:space': '0',
        'w:color': '000000'  # black color
    }
    
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        for key, value in border_attrs.items():
            border.set(qn(key), value)
        tblBorders.append(border)
    
    tblPr = tbl.tblPr
    existing_borders = tblPr.find(qn('w:tblBorders'))
    if existing_borders is not None:
        tblPr.remove(existing_borders)
    tblPr.append(tblBorders)

def set_row_height(table, height=200):
    """
    Sets the height of each row in the table.

    :param table: The table object from python-docx.
    :param height: Height in twips (1/20th of a point).
    """
    for row in table.rows:
        tr = row._tr
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), str(height))
        trHeight.set(qn('w:hRule'), 'exact')  # Use 'exact' to enforce row height

        trPr = tr.get_or_add_trPr()
        existing_trHeight = trPr.find(qn('w:trHeight'))
        if existing_trHeight is not None:
            trPr.remove(existing_trHeight)
        trPr.append(trHeight)
		
def concatenate_word_documents(doc_paths, save_path='Combined_Document.docx'):
    """
    Concatenates multiple Word documents into a single document, 
    separating contents from each document with an empty line.

    Parameters:
    - doc_paths: List of file paths to the Word documents to concatenate.
    - output_path: The file path to save the combined Word document.
    """
    # Create a new Document to store combined contents
    combined_doc = Document()

    for path in doc_paths:
        
        print(path)
        # Open each document
        doc = Document(path)
        
        # Add contents of the document to combined_doc
        for element in doc.element.body:
            combined_doc.element.body.append(element)
        
        # Append an empty paragraph to separate contents of different documents
        combined_doc.add_paragraph()

    # Save the combined document
    combined_doc.save(save_path)